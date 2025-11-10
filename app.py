# -*- coding: utf-8 -*-
import os
import json
from flask import Flask, render_template, request, send_file, jsonify, session, Response
from docxtpl import DocxTemplate
from datetime import datetime
from graph1 import run_graph, llm  # type: ignore
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Optional

app = Flask(__name__)
app.secret_key = "smart-rfp-ai-key"

from routes.table_routes import table_bp
app.register_blueprint(table_bp)


def fix_rtl_bullets(text: str) -> str:
    if not isinstance(text, str):
        return text
    rtl_start = "\u202B"
    rtl_end = "\u202C"
    replacements = {
        "•": f"{rtl_start}•{rtl_end}",
        "-": f"{rtl_start}-{rtl_end}",
        "–": f"{rtl_start}–{rtl_end}",
        "—": f"{rtl_start}—{rtl_end}",
        ":": f"{rtl_start}:{rtl_end}",
        "؛": f"{rtl_start}؛{rtl_end}",
        ".": f"{rtl_start}.{rtl_end}",
        "،": f"{rtl_start}،{rtl_end}",
        "(": f"{rtl_start}({rtl_end}",
        ")": f"{rtl_start}){rtl_end}",
        "[": f"{rtl_start}[{rtl_end}",
        "]": f"{rtl_start}]{rtl_end}",
        "{": f"{rtl_start}{{{rtl_end}",
        "}": f"{rtl_start}}}{rtl_end}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


# ----------------- ROUTES --------------------

@app.route('/')
def home() -> str:
    return render_template('home.html')


@app.route('/rfp_input')
def rfp_input() -> str:
    return render_template('rfp_input.html')


# ✅ NEW (GET / Streaming version)
@app.route('/rfp_generate', methods=['GET'])
def generate_stream():
    user_data = request.args.to_dict(flat=False)

    for key, value in user_data.items():
        if isinstance(value, list):
            user_data[key] = "، ".join(value)

    def generate_events():
        yield "data: جاري توليد الوثيقة...\n\n"

        result = run_graph(user_data)

        # ✅ Fix generator serialization issue
        if not isinstance(result, (dict, list, str, int, float, bool, type(None))):
            result = str(result)

        try:
            payload = json.dumps(result, ensure_ascii=False)
        except Exception:
            payload = json.dumps(str(result), ensure_ascii=False)

        yield f"data: {payload}\n\n"
        yield "data: {\"done\": true}\n\n"

    return Response(generate_events(), mimetype="text/event-stream")



# ✅ ORIGINAL POST (do NOT remove)
@app.route('/rfp_generate', methods=['POST'])
def generate():
    user_data = request.form.to_dict(flat=False)
    for key, value in user_data.items():
        if isinstance(value, list):
            user_data[key] = "، ".join(value)

    include_sections = {
        "Joint_Venture": request.form.get("include_Joint_Venture") is not None,
        "Tender_Split_Section": request.form.get("include_Tender_Split_Section") is not None,
        "Alternative_Offers": request.form.get("include_Alternative_Offers") is not None,
        "Insurance": request.form.get("include_Insurance") is not None,
    }
    session["include_sections"] = include_sections

    result = run_graph(user_data)
    decisions = result.get("decisions", {})

    if not decisions:
        return render_template("rfp_generate.html", decisions={}, user_data=user_data)

    from nodes.field_map import FIELD_MAP  # type: ignore
    filtered_decisions = {
        key: {"value": decisions.get(key, ""), "type": FIELD_MAP.get(key, "llm")}
        for key in FIELD_MAP
    }

    session["user_data"] = user_data
    session["decisions"] = {k: v["value"] for k, v in filtered_decisions.items()}

    return render_template("rfp_generate.html", decisions=filtered_decisions, user_data=user_data)


@app.route('/save', methods=['POST'])
def save():
    session_user = session.get("user_data", {})
    session_llm = session.get("decisions", {})
    edited_data = request.form.to_dict()

    context = {**session_user, **session_llm, **edited_data}
    tpl = DocxTemplate("templates/rfp_general.docx")

    TABLE_KEYS = {
        "Bill_of_Quantities_and_Prices",
        "Materials_Specifications_Table",
        "Equipment_Specifications_Table",
        "Workers_Table",
    }
    safe_context = {k: v for k, v in context.items() if k not in TABLE_KEYS}

    llm_fields = {
        "Competition_Definition",
        "Text_of_Costs_of_Competition_Documents",
        "Project_Scope_of_Work",
        "Technical_Proposal_Documents",
        "Financial_Proposal_Documents",
        "Bid_Evaluation_Criteria",
        "Regulatory_Records_and_Licenses",
        "Inquiries_Section",
        "Tender_Split_Section",
        "Penalties",
        "Delay_Penalties",
        "Insurance",
        "Service_Delivery_Plan",
        "Service_Execution_Method",
        "Alternative_Offers",
    }

    for key in list(safe_context.keys()):
        if key in llm_fields and isinstance(safe_context[key], str):
            safe_context[key] = fix_rtl_bullets(safe_context[key])

    tpl.render(safe_context)

    output_folder = os.path.join(os.path.expanduser("~"), "Documents", "RFP_outputs")
    os.makedirs(output_folder, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(output_folder, f"filled_{timestamp}.docx")
    tpl.save(output_path)

    def insert_table_if_exists(session_key, placeholder, heading_text=None):
        if session_key not in session:
            return

        table_text = session[session_key]
        lines = [l.strip() for l in table_text.split("\n") if "|" in l]
        if not lines:
            return

        headers = [h.strip() for h in lines[0].split("|")]
        rows = [l.split("|") for l in lines[1:]]

        doc = Document(output_path)

        def find_paragraph_with_text(doc_, text):
            for p in doc_.paragraphs:
                if text in p.text:
                    return p
            for table in doc_.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if text in p.text:
                                return p
            return None

        target_paragraph = None
        if heading_text:
            for p in doc.paragraphs:
                if p.text.strip() == heading_text.strip():
                    target_paragraph = p
                    break
        if not target_paragraph:
            target_paragraph = find_paragraph_with_text(doc, placeholder)
        if not target_paragraph:
            print(f"⚠️ لم يتم العثور على placeholder {placeholder}")
            return

        if placeholder in target_paragraph.text:
            target_paragraph.text = ""

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        tbl = table._element
        tbl.set(qn("w:tblDir"), "rtl")
        tbl.set(qn("w:tblLayout"), "fixed")

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for r in rows:
            row_cells = table.add_row().cells
            for i, c in enumerate(r):
                p = row_cells[i].paragraphs[0]
                run = p.add_run(c.strip())
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        target_paragraph._element.addnext(table._element)
        doc.save(output_path)

    insert_table_if_exists("Bill_of_Quantities_and_Prices", "{{Bill_of_Quantities_and_Prices}}", "جدول الكميات والأسعار")
    insert_table_if_exists("Materials_Specifications_Table", "{{Materials_Specifications_Table}}", "جدول مواصفات المواد")
    insert_table_if_exists("Equipment_Specifications_Table", "{{Equipment_Specifications_Table}}", "جدول مواصفات المعدات")
    insert_table_if_exists("Workers_Table", "{{Workers_Table}}", "جدول العمال")

    project_name = context.get("Competition_Name") or "مشروع بدون اسم"
    session["generated_file"] = output_path

    return render_template("result.html", project_name=project_name, download_url="/download")


@app.route('/download')
def download_file():
    file_path = session.get("generated_file")
    if not file_path or not os.path.exists(file_path):
        return "⚠️ الملف غير موجود. يرجى إعادة التوليد.", 404
    return send_file(file_path, as_attachment=True)
@app.route("/rfp_generate_done")
def rfp_generate_done():
    decisions = session.get("decisions", {})
    user_data = session.get("user_data", {})
    return render_template("rfp_generate.html", decisions=decisions, user_data=user_data)


if __name__ == "__main__":
    app.run(debug=True)
